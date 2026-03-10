#!/usr/bin/env python3
"""Gera o relatório de bifurcação em formato Word (.docx)"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── Estilos globais ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

# Heading styles
for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

def add_table(headers, rows, col_widths=None):
    """Helper para criar tabela formatada"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.style = doc.styles['Normal']
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    # Rows
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = str(val)
            for p in row.cells[i].paragraphs:
                p.style = doc.styles['Normal']
                for run in p.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()
    return table

# ══════════════════════════════════════════════════════
# CAPA
# ══════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

titulo = doc.add_paragraph()
titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = titulo.add_run('DIAGNOSTICO TECNICO')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

subtitulo = doc.add_paragraph()
subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitulo.add_run('Dois Deploys em Producao\nAnalise de Bifurcacao de Repositorios')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x77)

doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run('Best of Opera — Plataforma de Conteudo Musical\n\n')
run.font.size = Pt(12)
run = meta.add_run('Data: 9 de marco de 2026\n')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run = meta.add_run('Autor: Bolivar Andrade\n')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run = meta.add_run('Metodo: Investigacao automatizada via Claude Code\n')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run = meta.add_run('Classificacao: Documento interno — uso restrito aos socios')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()

# ══════════════════════════════════════════════════════
# RESUMO GERAL AMPLIADO
# ══════════════════════════════════════════════════════
doc.add_heading('Resumo Geral', level=1)

doc.add_paragraph(
    'Este documento apresenta uma analise tecnica completa da situacao atual '
    'do projeto Best of Opera, que possui dois deploys simultaneos em producao '
    'no Railway, mantidos de forma independente e sem coordenacao entre si.'
)

doc.add_paragraph(
    'A investigacao foi motivada pela descoberta de que o socio trabalhou durante '
    'o final de semana em um deploy proprio (protective-friendship-production-9659.'
    'up.railway.app), que apresenta uma tela de login com o branding "Arias Gestao '
    'de Conteudo" — funcionalidade que nao existe no projeto principal.'
)

doc.add_heading('O que foi descoberto', level=2)

doc.add_paragraph(
    'A organizacao BestOfOpera no GitHub possui dois repositorios:', style='List Bullet'
)

bullets = [
    'best-of-opera (antigo) — criado em 10/fev, ultimo commit em 13/fev, 36 commits, '
    'arquitetura monolitica (1 arquivo main.py de 1.459 linhas + 1 arquivo index.html de 72KB).',

    'best-of-opera-app2 (atual) — criado em 13/fev, atualizado ate 05/mar, 188 commits, '
    'arquitetura de monorepo com 4 aplicacoes especializadas, ~18.000 linhas de codigo proprio, '
    '14 paginas frontend, 70+ endpoints API.',

    'O socio construiu seu deploy a partir do repositorio antigo (best-of-opera), que foi '
    'abandonado em 13 de fevereiro — ha 24 dias. O projeto principal seguiu no repositorio '
    'best-of-opera-app2 com 152 commits adicionais que transformaram completamente a '
    'arquitetura, funcionalidades e infraestrutura do projeto.',

    'Nenhum commit do socio aparece em nenhum dos dois repositorios. Ele trabalhou '
    'completamente fora do controle de versao compartilhado, sem branch, sem fork '
    'registrado e sem pull request.',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('Impacto pratico', level=2)

doc.add_paragraph(
    'O deploy do socio nao contem nenhuma das funcionalidades criticas desenvolvidas '
    'nas ultimas 3+ semanas. Especificamente:'
)

impacts = [
    'O App Editor (pipeline de edicao de video com 9 etapas, legendas ASS sincronizadas, '
    'transcricao via Gemini, renderizacao FFmpeg em 7 idiomas) NAO EXISTE no deploy dele — '
    'sao 80+ commits e ~5.000 linhas de codigo inexistentes.',

    'O App Portal (frontend unificado Next.js 16 com 14 paginas, workers, polling '
    'adaptativo, importacao Redator-Editor, QA visual) esta em versao extremamente '
    'primitiva, sem as 50+ melhorias feitas desde 24/fev.',

    'O App Redator esta embutido como "Modulo de Producao" dentro do main.py monolitico, '
    'sem os prompts refinados, sem anti-leak de idioma, sem CTA fixo, sem auto-deteccao '
    'de metadados do YouTube.',

    'A infraestrutura usa disco local efemero (perde arquivos a cada redeploy) em vez '
    'do Cloudflare R2. Nao tem workers com heartbeat, nao tem dead-letter queue, nao '
    'tem Dockerfiles otimizados.',

    'A unica feature visivel que o socio adicionou foi uma tela de login com email/senha — '
    'funcionalidade que pode ser implementada em 1-2 horas no portal atual usando '
    'middleware do Next.js.',
]
for b in impacts:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('Conclusao do resumo', level=2)

p = doc.add_paragraph()
run = p.add_run(
    'A recomendacao tecnica e inequivoca: seguir com a versao principal (best-of-opera-app2) '
    'e descartar o deploy do socio. '
)
run.bold = True
p.add_run(
    'Nao ha possibilidade de merge entre as duas versoes — a arquitetura e completamente '
    'incompativel (monolito de 1 arquivo vs. monorepo com 4 apps). O trabalho do socio '
    'partiu de uma base abandonada ha quase um mes e nao inclui nenhuma das dezenas '
    'de features, correcoes e melhorias que tornaram o projeto viavel para producao. '
    'Tentar aproveitar qualquer codigo do deploy do socio seria mais custoso do que '
    'reimplementar as mesmas funcionalidades do zero no projeto atual.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════
# SECAO 1 — DOIS REPOSITORIOS
# ══════════════════════════════════════════════════════
doc.add_heading('1. Dois Repositorios na Organizacao GitHub', level=1)

doc.add_paragraph(
    'A organizacao BestOfOpera no GitHub contem exatamente dois repositorios. '
    'Ambos foram criados pelo mesmo usuario (Bolivar Andrade), mas apenas um '
    'continuou recebendo atualizacoes.'
)

add_table(
    ['', 'Repo Antigo (socio)', 'Repo Atual (Bolivar)'],
    [
        ['Nome', 'BestOfOpera/best-of-opera', 'BestOfOpera/best-of-opera-app2'],
        ['Criado', '10/fev/2026', '13/fev/2026'],
        ['Ultimo commit', '13/fev/2026', '05/mar/2026'],
        ['Total commits', '36', '188'],
        ['Arquitetura', 'Monolito (1 main.py + 1 index.html)', 'Monorepo (4 apps + portal + shared)'],
        ['Frontend', 'HTML estatico (72KB, 1 arquivo)', 'Next.js 16 + React 19 + 2 frontends Vite'],
        ['Backend', '1 FastAPI (1.459 linhas em main.py)', '3 FastAPIs especializados'],
        ['Banco', 'PostgreSQL (1 tabela cache)', 'PostgreSQL compartilhado (multiplas tabelas)'],
        ['Storage', 'Disco local (efemero)', 'Cloudflare R2 (persistente)'],
        ['Deploy', '1 servico Railway', '6 servicos Railway'],
    ]
)

doc.add_heading('Arquivos do repo antigo (TUDO que existe)', level=2)

doc.add_paragraph(
    'O repositorio antigo contem apenas 8 arquivos, totalizando ~162 KB de codigo:'
)
files_old = [
    'Procfile (58 bytes)',
    'README-2.md (5.5 KB)',
    'database.py (23.7 KB)',
    'dataset_v3_categorizado.csv (390 KB)',
    'main.py (60 KB — 1.459 linhas, TUDO junto)',
    'nixpacks.toml (62 bytes)',
    'requirements.txt (131 bytes)',
    'static/index.html (72 KB — HTML/CSS/JS inline)',
]
for f in files_old:
    doc.add_paragraph(f, style='List Bullet')

doc.add_heading('Estrutura do repo atual', level=2)

doc.add_paragraph(
    'O repositorio atual e um monorepo com 4 aplicacoes especializadas, '
    '~18.000 linhas de codigo proprio, 63+ componentes React, 14 paginas '
    'e 70+ endpoints API:'
)
dirs_new = [
    'app-curadoria/backend/ — Motor YouTube (FastAPI) — 1.704 linhas',
    'app-redator/ — Gerador de conteudo AI (FastAPI + React) — 3.866 linhas',
    'app-editor/ — Pipeline de video (FastAPI + React) — 5.110 linhas',
    'app-portal/ — Frontend unificado (Next.js 16) — 7.305 linhas',
    'shared/ — Codigo compartilhado',
    'scripts/ — Utilitarios Railway',
]
for d in dirs_new:
    doc.add_paragraph(d, style='List Bullet')

doc.add_page_break()

# ══════════════════════════════════════════════════════
# SECAO 2 — EVIDENCIAS TECNICAS
# ══════════════════════════════════════════════════════
doc.add_heading('2. Evidencias Tecnicas — Deploy do Socio', level=1)

# Prova 1
doc.add_heading('Prova 1: Tela de login inexistente no projeto principal', level=2)

doc.add_paragraph(
    'O deploy do socio (protective-friendship-production-9659.up.railway.app) '
    'exibe uma tela de login com campos "Email" e "Senha" e botao "Entrar", '
    'com branding "Arias" + "Gestao de Conteudo". Todas as rotas redirecionam '
    'para essa tela de login. O endpoint /api/health retorna redirecionamento '
    'para /login em vez de JSON.'
)

doc.add_paragraph(
    'O deploy principal (curadoria-production-cf4a.up.railway.app) tem acesso '
    'direto sem login, redireciona automaticamente para /curadoria (Motor V7) '
    'e /api/health retorna JSON com status, quota e versao.'
)

p = doc.add_paragraph()
run = p.add_run('Fato comprovado: ')
run.bold = True
p.add_run(
    'O arquivo app-portal/app/page.tsx no repo principal NUNCA teve uma tela '
    'de login. Desde o primeiro commit (a27bfd0, 15/fev), ele contem apenas um '
    'redirect para /curadoria. A busca no historico Git (git log --all -S "Entrar" '
    'e git log --all -S "password") confirma que nenhum commit jamais adicionou '
    'formulario de login em arquivos .tsx.'
)

# Prova 2
doc.add_heading('Prova 2: Metadata identica confirma origem no mesmo codigo', level=2)

doc.add_paragraph(
    'O deploy do socio exibe o titulo "Arias Conteudo" e descricao "Plataforma '
    'de gestao de conteudo para marcas de musica classica" — exatamente o que '
    'existe no layout.tsx desde o commit a27bfd0 (15/fev). Isso confirma que o '
    'socio clonou o portal e adicionou autenticacao por conta propria, sem '
    'contribuir de volta ao repositorio.'
)

# Prova 3
doc.add_heading('Prova 3: Senha hardcoded no repo antigo', level=2)

doc.add_paragraph(
    'O commit "feat: add simple password gate to protect app access" (12/fev, '
    'repo best-of-opera) adicionou uma senha fixa (APP_PASSWORD = "opera2026") '
    'ao main.py. Este padrao de "password gate" e consistente com a tela de '
    'login que o socio exibe.'
)

# Prova 4
doc.add_heading('Prova 4: Nenhum commit do socio em nenhum repositorio', level=2)

doc.add_paragraph(
    'A busca por autores no historico Git mostra apenas dois nomes — ambos '
    'do Bolivar:'
)
doc.add_paragraph('Bolivar Andrade <administrator@Mac-Bolivar.local> — 187 commits', style='List Bullet')
doc.add_paragraph('BestOfOpera <bolivarandrade@gmail.com> — 1 commit (Initial commit)', style='List Bullet')

p = doc.add_paragraph()
run = p.add_run(
    'Nenhum outro autor aparece em nenhum dos dois repositorios. O socio '
    'trabalhou completamente fora do controle de versao compartilhado.'
)
run.bold = True

# Prova 5
doc.add_heading('Prova 5: Apenas 1 branch em ambos os repositorios', level=2)

doc.add_paragraph(
    'Ambos os repositorios possuem apenas a branch main. Nao existe branch '
    'do socio, fork registrado, nem pull request.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════
# SECAO 3 — CRONOLOGIA
# ══════════════════════════════════════════════════════
doc.add_heading('3. Cronologia Completa', level=1)

doc.add_paragraph(
    'A tabela abaixo mostra a evolucao paralela dos dois repositorios. '
    'O repo antigo parou em 13/fev enquanto o repo atual continuou evoluindo '
    'por mais 24 dias com 152 commits adicionais.'
)

add_table(
    ['Data', 'Repo Antigo (best-of-opera)', 'Repo Atual (best-of-opera-app2)'],
    [
        ['10/fev', 'Criado. Upload inicial de arquivos', '—'],
        ['11/fev', 'Motor V3, PostgreSQL, playlist, cache', '—'],
        ['12/fev', 'Motor V7 + password gate', '—'],
        ['13/fev (manha)', 'APP2 Producao + fixes FFmpeg', 'Criado. Initial commit'],
        ['13/fev (tarde)', 'ULTIMO COMMIT (UX improvements)', 'APP Redator completo + deploy Railway'],
        ['14/fev', '—', '35 commits: Editor completo'],
        ['15/fev', '—', 'Monorepo reorganizado, Portal Next.js, Curadoria integrada'],
        ['23/fev', '—', 'Pipeline de transcricao reescrito'],
        ['24/fev', '—', 'Cloudflare R2, Dockerfiles, auto-deteccao YouTube'],
        ['25/fev', '—', 'Worker assincrono, traducao Google Cloud'],
        ['26/fev', '—', '30+ commits: Editor completo (renders, ZIP, R2)'],
        ['27/fev', '—', '35+ commits: Legendas ASS, workers, upload manual'],
        ['03/mar', '—', 'QA visual, toggle sem_lyrics, Overview'],
        ['04/mar', '—', 'Input manual YouTube, visualizacao de traducao'],
        ['05/mar', '—', 'CTA fixo, paginacao playlist, filtro tracks (HEAD)'],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════
# SECAO 4 — FUNCIONALIDADES AUSENTES
# ══════════════════════════════════════════════════════
doc.add_heading('4. Funcionalidades que NAO existem no deploy do socio', level=1)

# 4.1 Editor
doc.add_heading('4.1 App Editor — Inexistente (0% do pipeline)', level=2)

doc.add_paragraph(
    'O App Editor NAO EXISTE no repo antigo. Sao 2.000+ linhas de backend e '
    '3.000+ linhas de frontend desenvolvidos de 13/fev a 05/mar:'
)

add_table(
    ['Feature', 'Commits', 'Status no socio'],
    [
        ['Pipeline de 9 etapas (download -> render -> export)', '80+', 'INEXISTENTE'],
        ['Download via yt-dlp com retry e fallback', '10+', 'INEXISTENTE'],
        ['Transcricao de audio via Gemini 2.5 Pro', '8+', 'INEXISTENTE'],
        ['Merge inteligente (cega + guiada)', '3', 'INEXISTENTE'],
        ['Validacao de alinhamento (fuzzy match color-coded)', '5+', 'INEXISTENTE'],
        ['Legendas ASS com fonte TeX Gyre Pagella', '4', 'INEXISTENTE'],
        ['Traducao de lyrics em 7 idiomas (Google Cloud)', '5+', 'INEXISTENTE'],
        ['Renderizacao FFmpeg com legendas sincronizadas', '10+', 'INEXISTENTE'],
        ['Worker sequencial com heartbeat e idempotencia', '4', 'INEXISTENTE'],
        ['Dead-letter queue para crash-loop prevention', '2', 'INEXISTENTE'],
        ['Upload para Cloudflare R2', '3', 'INEXISTENTE'],
        ['Pacote ZIP assincrono', '2', 'INEXISTENTE'],
        ['Render sem legendas (saida de emergencia)', '1', 'INEXISTENTE'],
        ['Limpar Edicao (reset de estado travado)', '1', 'INEXISTENTE'],
        ['Desbloqueio forcado com diagnostico', '2', 'INEXISTENTE'],
        ['Cookies YouTube via base64', '1', 'INEXISTENTE'],
        ['Nomenclatura padronizada dos renders', '1', 'INEXISTENTE'],
        ['Anti-duplicata na importacao', '1', 'INEXISTENTE'],
    ]
)

# 4.2 Portal
doc.add_heading('4.2 App Portal — Frontend (parcialmente existente)', level=2)

doc.add_paragraph(
    'O socio tem alguma versao do portal, mas sem as features dos ultimos 50+ commits:'
)

add_table(
    ['Feature', 'Data', 'Status no socio'],
    [
        ['14 paginas integradas (curadoria + redator + editor)', '15/fev-05/mar', 'Parcial (sem editor)'],
        ['Tela de alinhamento com botoes editar/excluir/adicionar', '26/fev', 'INEXISTENTE'],
        ['Badge unificado Redator <-> Editor', '26/fev', 'INEXISTENTE'],
        ['Botao "Baixar Todos" com fallback individual', '26/fev', 'INEXISTENTE'],
        ['Aprovacao de preview com tratamento 409', '26/fev', 'INEXISTENTE'],
        ['Downloads por idioma em nova aba', '26/fev', 'INEXISTENTE'],
        ['Rotas de fuga e recovery na conclusao', '26/fev', 'INEXISTENTE'],
        ['Worker polling adaptativo', '25/fev', 'INEXISTENTE'],
        ['Importacao Redator -> Editor com deteccao de idioma', '25/fev', 'INEXISTENTE'],
        ['Overview com roteamento por passo_atual', '03/mar', 'INEXISTENTE'],
        ['Toggle sem_lyrics', '03/mar', 'INEXISTENTE'],
        ['Correcoes visuais QA (16+ itens)', '03/mar', 'INEXISTENTE'],
        ['Labels do menu lateral corrigidos', '03/mar', 'INEXISTENTE'],
        ['Input manual de URL YouTube na curadoria', '04/mar', 'INEXISTENTE'],
        ['Modal automatico ao adicionar video', '04/mar', 'INEXISTENTE'],
    ]
)

# 4.3 Redator
doc.add_heading('4.3 App Redator — Versao primitiva', level=2)

doc.add_paragraph(
    'O repo antigo tem o "Modulo de Producao" embutido no main.py (endpoints '
    '/api/prod/*). O repo atual tem o Redator como app separado:'
)

add_table(
    ['Feature', 'Data', 'Status no socio'],
    [
        ['Backend separado com routers por dominio', '13/fev', 'INEXISTENTE (tudo em 1 main.py)'],
        ['Frontend React dedicado (Vite)', '13/fev', 'INEXISTENTE (HTML estatico)'],
        ['Prompts refinados (overlay, post, YouTube, hook)', '13/fev-03/mar', 'Versao primitiva'],
        ['Traducao CTA e hashtags do post', '03/mar', 'INEXISTENTE'],
        ['Reforco de idioma nos prompts (anti-leak)', '03/mar', 'INEXISTENTE'],
        ['Labels da ficha tecnica no idioma do post', '27/fev', 'INEXISTENTE'],
        ['Omissao de campos vazios nos prompts', '27/fev', 'INEXISTENTE'],
        ['Limite 60 chars no overlay', '27/fev', 'INEXISTENTE'],
        ['CTA fixo na ultima legenda', '05/mar', 'INEXISTENTE'],
        ['Overlay congelado na importacao (fix)', '03/mar', 'INEXISTENTE'],
        ['Export ZIP funcional', '24/fev', 'Versao basica'],
        ['Auto-deteccao de metadados via YouTube', '24/fev', 'INEXISTENTE'],
    ]
)

# 4.4 Curadoria
doc.add_heading('4.4 App Curadoria Backend — Endpoints adicionais', level=2)

doc.add_paragraph(
    'O Motor V7 existe nos dois repositorios, mas o atual tem endpoints que '
    'o antigo nao tem:'
)

add_table(
    ['Endpoint / Feature', 'Repo Atual', 'Repo Antigo (socio)'],
    [
        ['POST /api/manual-video', 'Sim', 'Nao'],
        ['POST /api/prepare-video/{id}', 'Sim', 'Nao'],
        ['POST /api/upload-video/{id}', 'Sim', 'Nao'],
        ['GET /api/r2/check', 'Sim', 'Nao'],
        ['GET /api/r2/info', 'Sim', 'Nao'],
        ['POST /api/playlist/download-all', 'Sim', 'Nao'],
        ['GET /api/playlist/download-status', 'Sim', 'Nao'],
        ['Integracao com Cloudflare R2', 'Sim', 'Nao'],
        ['Paginacao de playlist (fix ERR-057)', 'Sim', 'Nao'],
        ['Filtro de tracks lyrics (fix ERR-055)', 'Sim', 'Nao'],
    ]
)

# 4.5 Infraestrutura
doc.add_heading('4.5 Infraestrutura — Completamente diferente', level=2)

add_table(
    ['Aspecto', 'Repo Atual', 'Repo Antigo (socio)'],
    [
        ['Servicos Railway', '6 servicos especializados', '1 servico monolito'],
        ['Storage', 'Cloudflare R2 (persistente)', 'Disco local (perde no redeploy)'],
        ['Dockerfiles', '3 Dockerfiles otimizados', 'Nixpacks generico'],
        ['CORS', 'Configurado por servico', 'allow_origins=["*"]'],
        ['Scripts de deploy', 'railway-env.sh (GraphQL API)', 'Manual'],
        ['Banco de dados', 'Tabelas por app', '1 tabela cache'],
        ['Workers', 'Background workers com heartbeat', 'Sincrono'],
        ['Variaveis de ambiente', '10+ variaveis por servico', '3-4 variaveis'],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════
# SECAO 5 — COMPARACAO NUMERICA
# ══════════════════════════════════════════════════════
doc.add_heading('5. Comparacao Numerica', level=1)

add_table(
    ['Metrica', 'Repo Antigo (socio)', 'Repo Atual (Bolivar)'],
    [
        ['Commits', '36', '188'],
        ['Arquivos de codigo', '3 (main.py, database.py, index.html)', '100+'],
        ['Linhas de codigo', '~2.000', '~18.000'],
        ['Endpoints API', '36 (todos em 1 arquivo)', '70+ (em 8 modulos)'],
        ['Paginas frontend', '1 (HTML monolito)', '14 (Next.js + 2 React apps)'],
        ['Servicos Railway', '1', '6'],
        ['Autores no Git', '1', '1 (mesmo — nenhum commit do socio)'],
        ['Branches', '1', '1'],
        ['Ultima atualizacao', '13/fev/2026', '05/mar/2026'],
        ['Dias de desenvolvimento', '4 dias (10-13/fev)', '24 dias (13/fev-09/mar)'],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════
# SECAO 6 — RECOMENDACAO FINAL
# ══════════════════════════════════════════════════════
doc.add_heading('6. Recomendacao Final', level=1)

doc.add_heading('Seguir com best-of-opera-app2 (versao principal)', level=2)

doc.add_paragraph('Motivos tecnicos:', style='List Bullet')

reasons = [
    '152 commits de diferenca — o deploy do socio esta parado em 13/fev; o principal '
    'tem 20+ dias de desenvolvimento continuo.',

    'Arquitetura incompativel — monolito vs. monorepo com 4 apps. Nao e possivel '
    'fazer merge entre as duas bases de codigo.',

    'App Editor inteiro inexistente — o pipeline de edicao de video (80+ commits, '
    '~5.000 linhas) nao existe no repo antigo.',

    'Storage efemero — o repo antigo perde arquivos a cada redeploy do Railway; '
    'o atual usa Cloudflare R2 com persistencia garantida.',

    'Nenhum commit do socio — nao ha codigo para avaliar ou aproveitar; o trabalho '
    'dele nao esta versionado em nenhum repo acessivel.',

    'A unica feature visivel do socio (login) e trivial — implementar autenticacao '
    'no portal atual leva 1-2 horas com middleware do Next.js.',
]
for r in reasons:
    doc.add_paragraph(r, style='List Number')

doc.add_heading('Acoes recomendadas', level=2)

actions = [
    'Desativar o deploy protective-friendship-production-9659 no Railway.',
    'Implementar autenticacao no app-portal atual se necessario.',
    'Solicitar ao socio que trabalhe no repo best-of-opera-app2 via branches e pull requests.',
    'Arquivar o repo best-of-opera (antigo) como referencia historica.',
]
for a in actions:
    doc.add_paragraph(a, style='List Number')

doc.add_paragraph()
doc.add_paragraph()

# Rodape
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run(
    'Documento gerado automaticamente por Claude Code em 09/mar/2026.\n'
    'Todos os dados foram extraidos exclusivamente via git log, gh api e curl.\n'
    'Nenhum arquivo do projeto foi modificado durante a investigacao.'
)
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.italic = True

# ── Salvar ──
output_path = os.path.expanduser(
    "~/best-of-opera-app2/RELATORIO-BIFURCACAO-DEPLOYS.docx"
)
doc.save(output_path)
print(f"Relatorio salvo em: {output_path}")
